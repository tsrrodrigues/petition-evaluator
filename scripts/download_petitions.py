#!/usr/bin/env python3
"""
Download DOCX files and extract text from petitions
"""
import json
import requests
from pathlib import Path
from docx import Document
import time

def download_file(url, output_path):
    """Download a file from URL"""
    response = requests.get(url, timeout=30)
    response.raise_for_status()
    
    with open(output_path, 'wb') as f:
        f.write(response.content)
    
    return True

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        print(f"Error extracting text: {e}")
        return None

def main():
    project_dir = Path(__file__).parent.parent
    data_dir = project_dir / 'data'
    petitions_dir = project_dir / 'petitions'
    petitions_dir.mkdir(exist_ok=True)
    
    # Load metadata
    metadata_file = data_dir / 'petitions_metadata.json'
    with open(metadata_file, 'r', encoding='utf-8') as f:
        petitions = json.load(f)
    
    print(f"Processing {len(petitions)} petitions...")
    
    results = []
    for i, petition in enumerate(petitions, 1):
        request_id = petition['request_id']
        rating = petition['rating']
        url = petition['url']
        
        print(f"\n[{i}/{len(petitions)}] Processing request_id={request_id}, rating={rating}")
        
        # Download DOCX
        docx_filename = f"{request_id}_rating{rating}.docx"
        docx_path = petitions_dir / docx_filename
        
        try:
            if not docx_path.exists():
                print(f"  Downloading from {url}...")
                download_file(url, docx_path)
                time.sleep(0.5)  # Be polite to the server
            else:
                print(f"  File already exists: {docx_filename}")
            
            # Extract text
            txt_filename = f"{request_id}_rating{rating}.txt"
            txt_path = petitions_dir / txt_filename
            
            if not txt_path.exists():
                print(f"  Extracting text...")
                text = extract_text_from_docx(docx_path)
                
                if text:
                    with open(txt_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    print(f"  ✓ Saved to {txt_filename} ({len(text)} chars)")
                else:
                    print(f"  ✗ Failed to extract text")
                    continue
            else:
                print(f"  Text file already exists: {txt_filename}")
                with open(txt_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            results.append({
                'request_id': request_id,
                'rating': rating,
                'docx_file': docx_filename,
                'txt_file': txt_filename,
                'text_length': len(text),
                'url': url,
                'remark': petition.get('remark'),
                'rating_text': petition.get('rating_text')
            })
            
        except Exception as e:
            print(f"  ✗ Error: {e}")
            continue
    
    # Save processing results
    results_file = data_dir / 'processed_petitions.json'
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*60}")
    print(f"Successfully processed {len(results)} out of {len(petitions)} petitions")
    print(f"Results saved to: {results_file}")
    
    # Summary by rating
    from collections import Counter
    rating_counts = Counter([r['rating'] for r in results])
    print("\nProcessed petitions by rating:")
    for rating in sorted(rating_counts.keys(), reverse=True):
        avg_length = sum(r['text_length'] for r in results if r['rating'] == rating) / rating_counts[rating]
        print(f"  Rating {rating}: {rating_counts[rating]} petitions (avg {avg_length:.0f} chars)")

if __name__ == '__main__':
    main()
